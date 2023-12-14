#import win32api
from prettytable import PrettyTable,ALL
from docx import Document
from docx.shared import Inches,Length,Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
#import win32api
def collection(file,document) :
 f=open(file)
 x1=f.read()
 f.close()
 x=x1.split('\n')
 values=['Region','Invoice No','Invoice Date','Retailer PAN']
 first,last,billval=[],[],[]
 bill=''
 invoice=[]
 name=[]
 f=open('config.txt')
 config=eval(f.read())
 f.close()
 
 for i in range(0,len(x)) :
    if 'Invoice No ' in x[i] and  config['secname'] in x[i] :
       first.append(i)
       invoice.append(x[i])
    if 'Time of Billing ' in x[i]  :
       last.append(i)
    if 'Bill Amount' in x[i] :
        billval.append(x[i])
    if 'Retailer ' in x[i] and 'Name' in x[i] and  config['secadd'] in x[i]:
        name.append(x[i])
        
 table=PrettyTable(['Date','Amount','Balance'])
 table.hrules=ALL
 table._min_height = 10
 table.width = 75
 style = document.styles['Normal']
 font = style.font
 font.name = 'Courier New'
 font.size = Pt(9)
 style.paragraph_format.space_before = Pt(0)
 style.paragraph_format.space_after = Pt(0)

 sections = document.sections
 for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

 for i in range(0,3) :
  table.add_row([' '*20]*3)
 for i in range(0,len(first)) :
    y1=x[first[i]:last[i]+1]
    #print(y1)
    for j in y1 :
        bill+=j+'\n'
        l=0
        for t in values :
         if t in j :
            l=1
            if 'Time' in j :
             j1=j
            else :
             j1=j.split(t)[0]
        if l==0 :
         paragraph=document.add_paragraph(j)
        else :
           paragraph=document.add_paragraph(j1) 
    billvalue=billval[i]
    billvalue1=billvalue.split('Bill')[0]
    billvalue2='Bill'+billvalue.split('Bill')[1]
    paragraph=document.add_paragraph(billvalue1)
    paragraph1=document.add_paragraph()
    print(invoice[i].split('Invoice')[1].split(':')[1])
    print(name)
    imp=invoice[i].split('Invoice')[1].split(':')[1]+'*'+name[i].split(':')[1]+'*'+'Amt :'+billvalue2.split(':')[1]
    imp=' '.join(imp.split())
    paragraph1=paragraph1.add_run('  '+'   '.join(imp.split('*')))
    paragraph1.font.size=Pt(12)
    paragraph1.bold=True
    paragraph3=document.add_paragraph().add_run(' '*60 +'Signature')
    paragraph3.alignment = 2
    paragraph3.font.size=Pt(12)
    paragraph3.bold=True
    f=open('config.txt')
    try : 
       lines = int(eval(f.read())['lines'])
    except : 
       lines=23
    f.close()
    if i%2 == 0:
     document.add_paragraph('\n'*lines) 
    else :
      #document.add_paragraph()  
      document.add_page_break()
    
      


def main(file,outputfile):
 document = Document()
 collection(file,document)
 document.save(outputfile)
 #win32api.ShellExecute (0,'print',file.split('.')[0]+'.docx',None, '.', 0 )
 
#main('','a.docx')
#main('recentmanual1.txt','recentmanual.docx')

